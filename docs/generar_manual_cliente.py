"""
Genera el "Manual de Usuario" en Word (.docx) para el cliente final.
Lenguaje sencillo, sin tecnicismos ni programación.

Uso:
    python docs/generar_manual_cliente.py
Crea/actualiza:  Manual de Usuario - Plataforma de Trading Automatico.docx
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches

AZUL = RGBColor(0x1F, 0x3B, 0x73)
GRIS = RGBColor(0x55, 0x55, 0x55)


def add_title(doc, texto, sub=None):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(texto)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = AZUL
    if sub:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = s.add_run(sub)
        r.font.size = Pt(14)
        r.font.color.rgb = GRIS


def h1(doc, texto):
    p = doc.add_heading(texto, level=1)
    for r in p.runs:
        r.font.color.rgb = AZUL


def h2(doc, texto):
    p = doc.add_heading(texto, level=2)
    for r in p.runs:
        r.font.color.rgb = AZUL


def parrafo(doc, texto):
    p = doc.add_paragraph(texto)
    p.paragraph_format.space_after = Pt(6)
    return p


def vinheta(doc, texto):
    doc.add_paragraph(texto, style="List Bullet")


def paso(doc, texto):
    doc.add_paragraph(texto, style="List Number")


def tabla(doc, encabezados, filas):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, e in enumerate(encabezados):
        cel = t.rows[0].cells[i]
        cel.text = ""
        run = cel.paragraphs[0].add_run(e)
        run.bold = True
    for fila in filas:
        celdas = t.add_row().cells
        for i, valor in enumerate(fila):
            celdas[i].text = valor
    doc.add_paragraph()
    return t


def construir():
    doc = Document()

    # Márgenes cómodos
    for s in doc.sections:
        s.top_margin = Inches(0.9)
        s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # Portada
    doc.add_paragraph("\n\n\n")
    add_title(doc, "Plataforma de Trading Automático",
              "Manual de Usuario")
    doc.add_paragraph("\n\n")
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ri = intro.add_run(
        "Tu cuenta operando sola, en la nube, con estrategias profesionales.\n"
        "Sin instalar nada. Sin dejar la computadora encendida."
    )
    ri.italic = True
    ri.font.color.rgb = GRIS
    doc.add_page_break()

    # 1
    h1(doc, "1. ¿Qué es esta plataforma?")
    parrafo(doc,
            "Es una herramienta que opera por ti en tu cuenta de trading. Tú "
            "eliges una estrategia y la plataforma se encarga de vigilar el "
            "mercado y abrir o cerrar operaciones automáticamente, las 24 horas.")
    parrafo(doc, "Sus ventajas principales:")
    vinheta(doc, "No necesitas instalar ningún programa en tu computadora.")
    vinheta(doc, "No tienes que dejar la PC encendida: todo funciona en la nube.")
    vinheta(doc, "Operas con estrategias ya programadas y probadas.")
    vinheta(doc, "Tú mantienes el control: activas, pausas o cambias cuando quieras.")

    # 2
    h1(doc, "2. ¿Cómo funciona por dentro? (en simple)")
    parrafo(doc,
            "No necesitas entender la parte técnica, pero esta es la idea en "
            "cuatro pasos:")
    paso(doc, "Conectas tu cuenta de broker a la plataforma.")
    paso(doc, "Eliges una estrategia y la configuras a tu gusto.")
    paso(doc, "La plataforma vigila el mercado en la nube de forma continua.")
    paso(doc, "Cuando se cumplen las condiciones de la estrategia, abre la "
              "operación por ti y le coloca su control de pérdidas y ganancias.")
    parrafo(doc,
            "Tú solo entras al panel cuando quieras para ver cómo va todo. "
            "Mientras tanto, el sistema trabaja por su cuenta.")

    # 3
    h1(doc, "3. Primeros pasos")
    paso(doc, "Entra a la plataforma con tu usuario y contraseña.")
    paso(doc, "Conecta tu cuenta de broker (ver sección 4).")
    paso(doc, "Crea tu primer robot y elige una estrategia (ver sección 5).")
    paso(doc, "Actívalo y supervísalo desde el panel.")
    parrafo(doc,
            "Consejo: la primera vez, usa una cuenta de práctica (demo). Así "
            "puedes ver todo funcionando sin arriesgar dinero real.")

    # 4
    h1(doc, "4. Conectar tu cuenta de broker")
    parrafo(doc,
            "En la sección «Cuentas» eliges «Conectar cuenta» y completas los "
            "datos que te dio tu broker:")
    vinheta(doc, "Un nombre para identificarla (ej. «Mi cuenta principal»).")
    vinheta(doc, "La plataforma (MT4 o MT5).")
    vinheta(doc, "El número de cuenta (login) y el servidor.")
    vinheta(doc, "La contraseña de la cuenta.")
    parrafo(doc,
            "Tu contraseña se usa solo para conectar la cuenta y no queda "
            "guardada en la plataforma. Tras conectarla, la cuenta tarda uno o "
            "dos minutos en quedar lista. Cuando aparezca como «lista», ya puedes "
            "usarla.")

    # 5
    h1(doc, "5. Crear tu primer robot")
    parrafo(doc,
            "Un «robot» (o bot) es una estrategia aplicada a una de tus cuentas. "
            "En la sección «Bots» eliges «Crear bot» y defines:")
    vinheta(doc, "Nombre del robot.")
    vinheta(doc, "La cuenta de broker donde va a operar.")
    vinheta(doc, "Los pares o símbolos (por ejemplo: EURUSD, XAUUSD).")
    vinheta(doc, "La estrategia que quieres usar (ver sección 6).")
    vinheta(doc, "El tamaño de la operación (lote) y tu control de riesgo.")
    parrafo(doc,
            "Cuando lo marcas como «activo» y guardas, el robot empieza a "
            "trabajar en pocos segundos.")

    # 6
    h1(doc, "6. Las estrategias disponibles")
    parrafo(doc, "Puedes elegir entre tres estrategias, explicadas en simple:")

    h2(doc, "Simple (dirección fija)")
    parrafo(doc,
            "Abre operaciones siempre en la dirección que tú elijas (compra o "
            "venta), respetando tus límites. Es la más básica, ideal para "
            "empezar o para acompañar una decisión manual.")

    h2(doc, "Multi-Tendencia con Flujo de Órdenes")
    parrafo(doc,
            "Solo opera cuando la tendencia apunta en la misma dirección en "
            "varios marcos de tiempo y, además, el movimiento del mercado lo "
            "confirma con fuerza. Es más selectiva: opera menos, pero busca "
            "entradas de mayor calidad.")

    h2(doc, "Ruptura del Rango Asiático")
    parrafo(doc,
            "Durante la madrugada (sesión asiática) el mercado suele moverse en "
            "un rango estrecho. Esta estrategia espera a que, ya en el horario "
            "de Londres o Nueva York, el precio rompa ese rango con fuerza, y "
            "entra en la dirección de la ruptura. Pensada para quienes buscan "
            "aprovechar los movimientos fuertes del día.")

    # 7
    h1(doc, "7. Modo de prueba: opera sin riesgo primero")
    parrafo(doc,
            "Antes de operar con dinero real, te recomendamos probar en una "
            "cuenta de práctica (demo). Algunas estrategias tienen además un "
            "«modo simulación», que muestra qué operación se habría hecho sin "
            "ejecutarla de verdad. Así puedes comprobar el comportamiento con "
            "total tranquilidad.")

    # 8
    h1(doc, "8. Supervisar tus operaciones")
    parrafo(doc, "Desde el panel puedes en todo momento:")
    vinheta(doc, "Ver cuántos robots tienes activos.")
    vinheta(doc, "Activar o pausar un robot con un clic.")
    vinheta(doc, "Cambiar la estrategia o los parámetros cuando quieras.")
    vinheta(doc, "Desconectar una cuenta.")
    parrafo(doc,
            "Las operaciones también aparecen en tu MetaTrader, como cualquier "
            "otra, identificadas con una etiqueta del robot.")

    # 9
    h1(doc, "9. Recomendaciones de seguridad")
    vinheta(doc, "Empieza siempre en cuenta demo o con un tamaño de operación pequeño.")
    vinheta(doc, "No inviertas dinero que no estés dispuesto a arriesgar.")
    vinheta(doc, "Revisa tu cuenta con regularidad, aunque el sistema sea automático.")
    vinheta(doc, "Mantén tu usuario y contraseña en privado.")
    parrafo(doc,
            "Recuerda: ninguna estrategia garantiza ganancias. El trading "
            "siempre implica riesgo de pérdida.")

    # 10
    h1(doc, "10. Preguntas frecuentes")
    tabla(doc, ["Pregunta", "Respuesta"], [
        ["¿Tengo que dejar la computadora encendida?",
         "No. Todo funciona en la nube, día y noche."],
        ["¿Puedo detener un robot cuando quiera?",
         "Sí, en cualquier momento desde el panel."],
        ["¿La plataforma garantiza ganancias?",
         "No. Ayuda a operar de forma automática, pero el mercado siempre tiene riesgo."],
        ["¿Mis datos del broker están seguros?",
         "La contraseña solo se usa para conectar y no queda guardada."],
        ["¿Puedo usar varias cuentas?",
         "Sí. Puedes conectar varias y poner robots distintos en cada una."],
    ])

    # 11
    h1(doc, "11. Glosario sencillo")
    tabla(doc, ["Término", "Qué significa"], [
        ["Broker", "La empresa donde tienes tu cuenta para operar."],
        ["Bot / Robot", "Una estrategia aplicada a una de tus cuentas."],
        ["Símbolo / Par", "El activo que se opera (ej. EURUSD, oro/XAUUSD)."],
        ["Lote", "El tamaño de la operación."],
        ["Stop Loss (SL)", "Límite de pérdida: cierra la operación si va en contra."],
        ["Take Profit (TP)", "Objetivo de ganancia: cierra la operación al alcanzarlo."],
        ["Demo", "Cuenta de práctica con dinero ficticio para probar sin riesgo."],
    ])

    doc.add_paragraph()
    cierre = doc.add_paragraph()
    cierre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cierre.add_run("Gracias por usar la plataforma. Opera con cabeza y disfruta el proceso.")
    rc.italic = True
    rc.font.color.rgb = GRIS

    salida = "Manual de Usuario - Plataforma de Trading Automatico.docx"
    doc.save(salida)
    print(f"Manual creado: {salida}")


if __name__ == "__main__":
    construir()
